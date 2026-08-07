# -*- coding: utf-8 -*-
from pathlib import Path
import sys

from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

DOCS_SCRIPTS = Path(r"C:\Users\thean\.codex\plugins\cache\openai-primary-runtime\documents\26.623.12021\skills\documents\scripts")
if str(DOCS_SCRIPTS) not in sys.path:
    sys.path.append(str(DOCS_SCRIPTS))

try:
    from table_geometry import apply_table_geometry
except ImportError:
    def apply_table_geometry(table, widths_dxa, **kwargs):
        pass

ROOT = Path(__file__).resolve().parents[1]
OUT_PATH = ROOT / "Baocaotiendotuan5_NguyenTheAn_2400004657.docx"

FONT_NAME = "Times New Roman"
BODY_SIZE = 13
TITLE_SIZE = 16
HEADING_SIZE = 14
BLUE = RGBColor(15, 71, 97)
BLACK = RGBColor(0, 0, 0)

def set_run_font(run, *, size=None, color=None, bold=None, italic=None):
    run.font.name = FONT_NAME
    run._element.rPr.rFonts.set(qn("w:ascii"), FONT_NAME)
    run._element.rPr.rFonts.set(qn("w:hAnsi"), FONT_NAME)
    run._element.rPr.rFonts.set(qn("w:eastAsia"), FONT_NAME)
    if size is not None:
        run.font.size = Pt(size)
    if color is not None:
        run.font.color.rgb = color
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic

def set_style_font(style, *, size=None, color=None, bold=None, italic=None):
    style.font.name = FONT_NAME
    style.element.rPr.rFonts.set(qn("w:ascii"), FONT_NAME)
    style.element.rPr.rFonts.set(qn("w:hAnsi"), FONT_NAME)
    style.element.rPr.rFonts.set(qn("w:eastAsia"), FONT_NAME)
    if size is not None:
        style.font.size = Pt(size)
    if color is not None:
        style.font.color.rgb = color
    if bold is not None:
        style.font.bold = bold
    if italic is not None:
        style.font.italic = italic

def format_paragraph(paragraph, *, alignment=WD_ALIGN_PARAGRAPH.JUSTIFY, before=0, after=0, line=1.5, first_line=0):
    paragraph.alignment = alignment
    fmt = paragraph.paragraph_format
    fmt.space_before = Pt(before)
    fmt.space_after = Pt(after)
    fmt.line_spacing = line
    if first_line:
        fmt.first_line_indent = Inches(first_line)

def add_text_paragraph(doc, text, *, style="Normal", alignment=WD_ALIGN_PARAGRAPH.JUSTIFY, before=0, after=0, line=1.5, first_line=0, size=BODY_SIZE, color=BLACK, bold=False, italic=False):
    p = doc.add_paragraph(style=style)
    format_paragraph(p, alignment=alignment, before=before, after=after, line=line, first_line=first_line)
    run = p.add_run(text)
    set_run_font(run, size=size, color=color, bold=bold, italic=italic)
    return p

def clear_cell(cell):
    cell.text = ""
    for p in cell.paragraphs:
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(0)
        p.paragraph_format.line_spacing = 1.0

def set_cell_text(cell, text, *, bold=False, align=WD_ALIGN_PARAGRAPH.LEFT, size=BODY_SIZE):
    clear_cell(cell)
    p = cell.paragraphs[0]
    p.alignment = align
    fmt = p.paragraph_format
    fmt.space_before = Pt(0)
    fmt.space_after = Pt(0)
    fmt.line_spacing = 1.15
    run = p.add_run(text)
    set_run_font(run, size=size, bold=bold, color=BLACK)

def style_table(table, widths_dxa):
    apply_table_geometry(
        table,
        widths_dxa,
        table_width_dxa=sum(widths_dxa),
        indent_dxa=120,
        cell_margins_dxa={"top": 80, "bottom": 80, "start": 120, "end": 120},
    )
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    for row in table.rows:
        for cell in row.cells:
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.TOP
            for p in cell.paragraphs:
                p.paragraph_format.space_before = Pt(0)
                p.paragraph_format.space_after = Pt(0)
                p.paragraph_format.line_spacing = 1.15

def build_report():
    doc = Document()
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)

    normal = doc.styles["Normal"]
    set_style_font(normal, size=BODY_SIZE, color=BLACK)
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(0)
    normal.paragraph_format.line_spacing = 1.5

    h2 = doc.styles["Heading 2"]
    set_style_font(h2, size=TITLE_SIZE, color=BLUE, bold=True)
    h2.paragraph_format.space_before = Pt(0)
    h2.paragraph_format.space_after = Pt(6)
    h2.paragraph_format.line_spacing = 1.0

    h3 = doc.styles["Heading 3"]
    set_style_font(h3, size=HEADING_SIZE, color=BLUE, bold=True)
    h3.paragraph_format.space_before = Pt(12)
    h3.paragraph_format.space_after = Pt(6)
    h3.paragraph_format.line_spacing = 1.0

    doc.core_properties.title = "Báo cáo tiến độ tuần 5"
    doc.core_properties.author = "Nguyễn Thế An"
    doc.core_properties.subject = "PandoraWeb"

    add_text_paragraph(doc, "BÁO CÁO TIẾN ĐỘ TUẦN 5", style="Heading 2", alignment=WD_ALIGN_PARAGRAPH.LEFT, after=6, line=1.0, size=TITLE_SIZE, color=BLUE, bold=True)
    add_text_paragraph(doc, "Đề tài: Xây dựng hệ thống thương mại điện tử quản lý và kinh doanh trang sức trực tuyến", style="Normal", alignment=WD_ALIGN_PARAGRAPH.JUSTIFY, after=8, line=1.5, size=BODY_SIZE, color=BLACK)

    meta = doc.add_table(rows=3, cols=2)
    style_table(meta, [2940, 6420])
    set_cell_text(meta.cell(0, 0), "Sinh viên thực hiện", bold=True)
    set_cell_text(meta.cell(0, 1), "Nguyễn Thế An")
    set_cell_text(meta.cell(1, 0), "MSSV", bold=True)
    set_cell_text(meta.cell(1, 1), "2400004657")
    set_cell_text(meta.cell(2, 0), "Thời gian thực hiện", bold=True)
    set_cell_text(meta.cell(2, 1), "Tuần 5")

    add_text_paragraph(doc, "1. Tổng quan tiến độ tuần 5:", style="Heading 3", alignment=WD_ALIGN_PARAGRAPH.LEFT, after=6, line=1.0, size=HEADING_SIZE, color=BLUE, bold=True)
    add_text_paragraph(doc, "Sau khi hoàn thiện module Giỏ hàng (Shopping Cart) và lưu trữ tạm thời qua Session ở tuần 4, tuần 5 của dự án PandoraWeb được tập trung vào hoàn thiện phân hệ Đặt hàng và Thanh toán (Checkout) – xử lý nghiệp vụ đơn hàng sâu dưới Database. Trong tuần này, em đã xây dựng trang Thanh toán nhận thông tin giao hàng của khách, xử lý luồng tạo tự động Khách hàng vãng lai (Guest Checkout), chuyển đổi dữ liệu từ Giỏ hàng Session lưu vĩnh viễn xuống các bảng Orders & OrderItems trong CSDL, đồng thời phát triển thuật toán tự động trừ tồn kho và gửi email hóa đơn xác nhận ngay khi đặt hàng thành công.", style="Normal", alignment=WD_ALIGN_PARAGRAPH.JUSTIFY, after=6, line=1.5, first_line=0.5, size=BODY_SIZE, color=BLACK)

    add_text_paragraph(doc, "2. Công việc đã thực hiện trong tuần 5:", style="Heading 3", alignment=WD_ALIGN_PARAGRAPH.LEFT, after=6, line=1.0, size=HEADING_SIZE, color=BLUE, bold=True)
    add_text_paragraph(doc, "Các đầu việc trong tuần 5 được triển khai bám sát 4 yêu cầu nghiệp vụ cốt lõi của phân hệ đặt hàng và thanh toán:", style="Normal", alignment=WD_ALIGN_PARAGRAPH.JUSTIFY, after=6, line=1.5, first_line=0.5, size=BODY_SIZE, color=BLACK)

    work = doc.add_table(rows=5, cols=3)
    style_table(work, [850, 2500, 6010])
    set_cell_text(work.cell(0, 0), "STT", bold=True, align=WD_ALIGN_PARAGRAPH.CENTER)
    set_cell_text(work.cell(0, 1), "Hạng mục", bold=True, align=WD_ALIGN_PARAGRAPH.CENTER)
    set_cell_text(work.cell(0, 2), "Nội dung thực hiện", bold=True, align=WD_ALIGN_PARAGRAPH.CENTER)

    set_cell_text(work.cell(1, 0), "1", align=WD_ALIGN_PARAGRAPH.CENTER)
    set_cell_text(work.cell(1, 1), "Tích hợp trang Thanh toán nhận thông tin giao hàng")
    set_cell_text(work.cell(1, 2), "Xây dựng giao diện và Controller (Checkout trong OrderController.cs) thu thập thông tin nhận hàng của khách (Họ tên, SĐT, Email, Địa chỉ chi tiết, Ghi chú). Tự động điền sẵn thông tin nếu khách đã đăng nhập. Hỗ trợ 2 phương thức thanh toán: Thanh toán khi nhận hàng (COD) và Chuyển khoản ngân hàng (BANK / QR Code).")

    set_cell_text(work.cell(2, 0), "2", align=WD_ALIGN_PARAGRAPH.CENTER)
    set_cell_text(work.cell(2, 1), "Tự động tạo Khách hàng vãng lai (Guest Checkout)")
    set_cell_text(work.cell(2, 2), "Xây dựng luồng xử lý kiểm tra email khi chốt đơn. Nếu người mua chưa có tài khoản trong hệ thống, Controller tự động khởi tạo hồ sơ Khách hàng vãng lai (Customer) mới trong Database với trạng thái active, đồng thời tự động đăng nhập duy trì Session cho khách để người dùng xem ngay lịch sử đơn hàng mà không làm gián đoạn trải nghiệm.")

    set_cell_text(work.cell(3, 0), "3", align=WD_ALIGN_PARAGRAPH.CENTER)
    set_cell_text(work.cell(3, 1), "Chuyển dữ liệu Giỏ hàng (Session) lưu vĩnh viễn xuống Database")
    set_cell_text(work.cell(3, 2), "Viết hàm chuyển đổi toàn bộ thông tin giỏ hàng tạm thời từ Session[\"Cart\"] thành các bản ghi vĩnh viễn trong 2 bảng Orders và OrderItems. Sau khi lưu thành công, hệ thống xóa sạch giỏ hàng tạm trong Session/Database, đồng thời tự động kích hoạt dịch vụ EmailHelper gửi Email hóa đơn xác nhận chi tiết về Gmail người mua.")

    set_cell_text(work.cell(4, 0), "4", align=WD_ALIGN_PARAGRAPH.CENTER)
    set_cell_text(work.cell(4, 1), "Thuật toán Tự động trừ tồn kho khi đặt hàng thành công")
    set_cell_text(work.cell(4, 2), "Áp dụng Database Transaction (db.Database.BeginTransaction()) kiểm tra số lượng tồn kho của từng biến thể sản phẩm trong bảng ProductVariants. Nếu đủ hàng, hệ thống thực hiện trừ tồn kho (Stock -= Quantity) và chốt giao dịch. Nếu không đủ, tự động Rollback hoàn tác toàn bộ và hiển thị thông báo lỗi rõ ràng.")

    add_text_paragraph(doc, "3. Kết quả đạt được:", style="Heading 3", alignment=WD_ALIGN_PARAGRAPH.LEFT, after=6, line=1.0, size=HEADING_SIZE, color=BLUE, bold=True)
    add_text_paragraph(doc, "Đến cuối tuần 5, quy trình mua sắm trên PandoraWeb đã khép kín hoàn chỉnh từ bước duyệt hàng, cho vào giỏ, điền thông tin thanh toán, lưu đơn vĩnh viễn xuống Database, trừ tồn kho chính xác cho đến gửi email xác nhận. Hệ thống phục vụ mượt mà cho cả người dùng đăng nhập sẵn lẫn khách mua nhanh vãng lai.", style="Normal", alignment=WD_ALIGN_PARAGRAPH.JUSTIFY, after=6, line=1.5, first_line=0.5, size=BODY_SIZE, color=BLACK)
    add_text_paragraph(doc, "So với tuần 4, ứng dụng đã tiến một bước lớn từ quản lý giỏ hàng tạm thời sang việc xử lý giao dịch thương mại điện tử thực sự với Database và Email dịch vụ.", style="Normal", alignment=WD_ALIGN_PARAGRAPH.JUSTIFY, after=6, line=1.5, first_line=0.5, size=BODY_SIZE, color=BLACK)

    add_text_paragraph(doc, "4. Khó khăn và hướng xử lý:", style="Heading 3", alignment=WD_ALIGN_PARAGRAPH.LEFT, after=6, line=1.0, size=HEADING_SIZE, color=BLUE, bold=True)

    diff = doc.add_table(rows=4, cols=3)
    style_table(diff, [850, 2600, 5910])
    set_cell_text(diff.cell(0, 0), "STT", bold=True, align=WD_ALIGN_PARAGRAPH.CENTER)
    set_cell_text(diff.cell(0, 1), "Khó khăn", bold=True, align=WD_ALIGN_PARAGRAPH.CENTER)
    set_cell_text(diff.cell(0, 2), "Hướng xử lý", bold=True, align=WD_ALIGN_PARAGRAPH.CENTER)

    set_cell_text(diff.cell(1, 0), "1", align=WD_ALIGN_PARAGRAPH.CENTER)
    set_cell_text(diff.cell(1, 1), "Tranh chấp dữ liệu tồn kho khi nhiều người mua cùng lúc")
    set_cell_text(diff.cell(1, 2), "Sử dụng Database Transaction (BeginTransaction) khóa và re-verify số lượng kho của ProductVariants ngay trước khi trừ, tự động Rollback nếu không đủ hàng.")

    set_cell_text(diff.cell(2, 0), "2", align=WD_ALIGN_PARAGRAPH.CENTER)
    set_cell_text(diff.cell(2, 1), "Đảm bảo trải nghiệm mua hàng nhanh cho khách chưa có tài khoản")
    set_cell_text(diff.cell(2, 2), "Tạo tài khoản Guest tự động dựa trên Email giao hàng, cấp Session truy cập ngay lập tức giúp khách hàng theo dõi đơn hàng dễ dàng.")

    set_cell_text(diff.cell(3, 0), "3", align=WD_ALIGN_PARAGRAPH.CENTER)
    set_cell_text(diff.cell(3, 1), "Gửi email xác nhận bị trễ hoặc chậm tiến trình chuyển trang")
    set_cell_text(diff.cell(3, 2), "Đưa tác vụ gửi mail vào luồng chạy ngầm bất đồng bộ (Task.Run), đồng thời lưu bản sao HTML hóa đơn tại server để kiểm tra nhanh.")

    add_text_paragraph(doc, "5. Kế hoạch tuần 6:", style="Heading 3", alignment=WD_ALIGN_PARAGRAPH.LEFT, after=6, line=1.0, size=HEADING_SIZE, color=BLUE, bold=True)
    add_text_paragraph(doc, "Trong tuần 6, em sẽ tập trung hoàn thiện phân hệ Bảo mật, Phân quyền và Hồ sơ cá nhân người dùng, cụ thể:", style="Normal", alignment=WD_ALIGN_PARAGRAPH.JUSTIFY, after=4, line=1.5, first_line=0.5, size=BODY_SIZE, color=BLACK)
    
    add_text_paragraph(doc, "Biến form Đăng ký và Đăng nhập tĩnh thành form động, thực hiện mã hóa mật khẩu an toàn xuống Database.", style="Normal", alignment=WD_ALIGN_PARAGRAPH.JUSTIFY, after=3, line=1.5, first_line=0.5, size=BODY_SIZE, color=BLACK)
    add_text_paragraph(doc, "Áp dụng bộ lọc bảo vệ ([AdminAuthorize]) để chặn triệt để người dùng không có quyền truy cập vào các trang Quản trị Admin.", style="Normal", alignment=WD_ALIGN_PARAGRAPH.JUSTIFY, after=3, line=1.5, first_line=0.5, size=BODY_SIZE, color=BLACK)
    add_text_paragraph(doc, "Hoàn thiện trang Hồ sơ cá nhân: Cho phép người dùng đổi mật khẩu, cập nhật thông tin cá nhân/ảnh đại diện, xem lại danh sách đơn hàng đã đặt và theo dõi chi tiết trạng thái giao hàng.", style="Normal", alignment=WD_ALIGN_PARAGRAPH.JUSTIFY, after=6, line=1.5, first_line=0.5, size=BODY_SIZE, color=BLACK)
    add_text_paragraph(doc, "Ngoài ra, em sẽ tiến hành rà soát tổng thể giao diện, kiểm thử bảo mật và chuẩn bị dữ liệu báo cáo cho giai đoạn nghiệm thu cuối cùng.", style="Normal", alignment=WD_ALIGN_PARAGRAPH.JUSTIFY, after=6, line=1.5, first_line=0.5, size=BODY_SIZE, color=BLACK)

    add_text_paragraph(doc, "6. Kết luận", style="Heading 3", alignment=WD_ALIGN_PARAGRAPH.LEFT, after=6, line=1.0, size=HEADING_SIZE, color=BLUE, bold=True)
    add_text_paragraph(doc, "Tuần 5 đã giải quyết xong toàn bộ phần cốt lõi của nghiệp vụ Đặt hàng - Thanh toán - Trừ tồn kho - Gửi hóa đơn Email. Việc kết nối chặt chẽ giữa Session, Database Transaction và Email Helper giúp hệ thống vận hành ổn định, chính xác và chuyên nghiệp theo chuẩn một website E-commerce thực thụ.", style="Normal", alignment=WD_ALIGN_PARAGRAPH.JUSTIFY, after=4, line=1.5, first_line=0.5, size=BODY_SIZE, color=BLACK)

    add_text_paragraph(doc, "Link tiến độ: https://docs.google.com/spreadsheets/d/1883ETMajLVA1eyZFIwByjCHXHZjkqqK9pXgpOd2hDI/edit?usp=sharing", style="Normal", alignment=WD_ALIGN_PARAGRAPH.JUSTIFY, after=0, line=1.5, first_line=0, size=BODY_SIZE, color=BLACK)

    doc.save(OUT_PATH)
    return OUT_PATH

if __name__ == "__main__":
    path = build_report()
    print(f"Saved week 5 report to {path}")
