# -*- coding: utf-8 -*-
from pathlib import Path
import sys

from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


DOCS_SCRIPTS = Path(r"C:\Users\thean\.codex\plugins\cache\openai-primary-runtime\documents\26.623.12021\skills\documents\scripts")
if str(DOCS_SCRIPTS) not in sys.path:
    sys.path.append(str(DOCS_SCRIPTS))

from table_geometry import apply_table_geometry


ROOT = Path(__file__).resolve().parents[1]
OUT_PATH = ROOT / "Baocaotiendotuan3_NguyenTheAn_2400004657.docx"

FONT_NAME = "Times New Roman"
BODY_SIZE = 13
TITLE_SIZE = 16
HEADING_SIZE = 14
CAPTION_SIZE = 9
BLUE = RGBColor(15, 71, 97)
BLACK = RGBColor(0, 0, 0)


def set_run_font(run, *, size=None, color=None, bold=None, italic=None):
    run.font.name = FONT_NAME
    run._element.rPr.rFonts.set(qn("w:ascii"), FONT_NAME)
    run._element.rPr.rFonts.set(qn("w:hAnsi"), FONT_NAME)
    run._element.rPr.rFonts.set(qn("w:eastAsia"), FONT_NAME)
    if size is not None:
        run.font.size = Pt(size)
    if color is not None:
        run.font.color.rgb = color
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic


def set_style_font(style, *, size=None, color=None, bold=None, italic=None):
    style.font.name = FONT_NAME
    style.element.rPr.rFonts.set(qn("w:ascii"), FONT_NAME)
    style.element.rPr.rFonts.set(qn("w:hAnsi"), FONT_NAME)
    style.element.rPr.rFonts.set(qn("w:eastAsia"), FONT_NAME)
    if size is not None:
        style.font.size = Pt(size)
    if color is not None:
        style.font.color.rgb = color
    if bold is not None:
        style.font.bold = bold
    if italic is not None:
        style.font.italic = italic


def format_paragraph(paragraph, *, alignment=WD_ALIGN_PARAGRAPH.JUSTIFY, before=0, after=0, line=1.5, first_line=0):
    paragraph.alignment = alignment
    fmt = paragraph.paragraph_format
    fmt.space_before = Pt(before)
    fmt.space_after = Pt(after)
    fmt.line_spacing = line
    if first_line:
        fmt.first_line_indent = Inches(first_line)


def add_text_paragraph(
    doc,
    text,
    *,
    style="Normal",
    alignment=WD_ALIGN_PARAGRAPH.JUSTIFY,
    before=0,
    after=0,
    line=1.5,
    first_line=0,
    size=BODY_SIZE,
    color=BLACK,
    bold=False,
    italic=False,
):
    p = doc.add_paragraph(style=style)
    format_paragraph(p, alignment=alignment, before=before, after=after, line=line, first_line=first_line)
    run = p.add_run(text)
    set_run_font(run, size=size, color=color, bold=bold, italic=italic)
    return p


def clear_cell(cell):
    cell.text = ""
    for p in cell.paragraphs:
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(0)
        p.paragraph_format.line_spacing = 1.0


def set_cell_text(cell, text, *, bold=False, align=WD_ALIGN_PARAGRAPH.LEFT, size=BODY_SIZE):
    clear_cell(cell)
    p = cell.paragraphs[0]
    p.alignment = align
    fmt = p.paragraph_format
    fmt.space_before = Pt(0)
    fmt.space_after = Pt(0)
    fmt.line_spacing = 1.15
    run = p.add_run(text)
    set_run_font(run, size=size, bold=bold, color=BLACK)


def style_table(table, widths_dxa):
    apply_table_geometry(
        table,
        widths_dxa,
        table_width_dxa=sum(widths_dxa),
        indent_dxa=120,
        cell_margins_dxa={"top": 80, "bottom": 80, "start": 120, "end": 120},
    )
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    for row in table.rows:
        for cell in row.cells:
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.TOP
            for p in cell.paragraphs:
                p.paragraph_format.space_before = Pt(0)
                p.paragraph_format.space_after = Pt(0)
                p.paragraph_format.line_spacing = 1.15


def build_report():
    doc = Document()
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    normal = doc.styles["Normal"]
    set_style_font(normal, size=BODY_SIZE, color=BLACK)
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(0)
    normal.paragraph_format.line_spacing = 1.5

    h2 = doc.styles["Heading 2"]
    set_style_font(h2, size=TITLE_SIZE, color=BLUE, bold=True)
    h2.paragraph_format.space_before = Pt(0)
    h2.paragraph_format.space_after = Pt(6)
    h2.paragraph_format.line_spacing = 1.0

    h3 = doc.styles["Heading 3"]
    set_style_font(h3, size=HEADING_SIZE, color=BLUE, bold=True)
    h3.paragraph_format.space_before = Pt(12)
    h3.paragraph_format.space_after = Pt(6)
    h3.paragraph_format.line_spacing = 1.0

    doc.core_properties.title = "Báo cáo tiến độ tuần 3"
    doc.core_properties.author = "Nguyễn Thế An"
    doc.core_properties.subject = "PandoraWeb"

    add_text_paragraph(
        doc,
        "BÁO CÁO TIẾN ĐỘ TUẦN 3",
        style="Heading 2",
        alignment=WD_ALIGN_PARAGRAPH.LEFT,
        after=6,
        line=1.0,
        size=TITLE_SIZE,
        color=BLUE,
        bold=True,
    )
    add_text_paragraph(
        doc,
        "Đề tài: Xây dựng hệ thống thương mại điện tử quản lý và kinh doanh trang sức trực tuyến",
        style="Normal",
        alignment=WD_ALIGN_PARAGRAPH.JUSTIFY,
        after=8,
        line=1.5,
        first_line=0,
        size=BODY_SIZE,
        color=BLACK,
    )

    meta = doc.add_table(rows=3, cols=2)
    style_table(meta, [2940, 6420])
    set_cell_text(meta.cell(0, 0), "Sinh viên thực hiện", bold=True)
    set_cell_text(meta.cell(0, 1), "Nguyễn Thế An")
    set_cell_text(meta.cell(1, 0), "MSSV", bold=True)
    set_cell_text(meta.cell(1, 1), "2400004657")
    set_cell_text(meta.cell(2, 0), "Thời gian thực hiện", bold=True)
    set_cell_text(meta.cell(2, 1), "Tuần 3")

    add_text_paragraph(
        doc,
        "1. Tổng quan tiến độ tuần 3:",
        style="Heading 3",
        alignment=WD_ALIGN_PARAGRAPH.LEFT,
        after=6,
        line=1.0,
        size=HEADING_SIZE,
        color=BLUE,
        bold=True,
    )
    add_text_paragraph(
        doc,
        "Sau khi hoàn tất nền tảng dữ liệu và các controller cốt lõi ở tuần 2, tuần 3 của dự án PandoraWeb được chuyển sang giai đoạn hoàn thiện luồng trải nghiệm người dùng và ghép nối các chức năng mua sắm quan trọng. Trong tuần này, em tập trung vào việc làm cho website không chỉ hiển thị dữ liệu động mà còn có thể phục vụ đầy đủ các bước từ xem sản phẩm, thêm vào giỏ hàng, đăng nhập, thanh toán cho đến quản trị dữ liệu ở phía sau.",
        style="Normal",
        alignment=WD_ALIGN_PARAGRAPH.JUSTIFY,
        after=6,
        line=1.5,
        first_line=0.5,
        size=BODY_SIZE,
        color=BLACK,
    )

    add_text_paragraph(
        doc,
        "2. Công việc đã thực hiện trong tuần 3:",
        style="Heading 3",
        alignment=WD_ALIGN_PARAGRAPH.LEFT,
        after=6,
        line=1.0,
        size=HEADING_SIZE,
        color=BLUE,
        bold=True,
    )
    add_text_paragraph(
        doc,
        "Các đầu việc trong tuần 3 được triển khai theo hướng hoàn thiện trải nghiệm end-to-end của website: từ trang chủ, bộ sưu tập sản phẩm, chi tiết sản phẩm cho đến giỏ hàng, thanh toán và khu vực quản trị.",
        style="Normal",
        alignment=WD_ALIGN_PARAGRAPH.JUSTIFY,
        after=6,
        line=1.5,
        first_line=0.5,
        size=BODY_SIZE,
        color=BLACK,
    )

    work = doc.add_table(rows=5, cols=3)
    style_table(work, [850, 2500, 6010])
    set_cell_text(work.cell(0, 0), "STT", bold=True, align=WD_ALIGN_PARAGRAPH.CENTER)
    set_cell_text(work.cell(0, 1), "Hạng mục", bold=True, align=WD_ALIGN_PARAGRAPH.CENTER)
    set_cell_text(work.cell(0, 2), "Nội dung thực hiện", bold=True, align=WD_ALIGN_PARAGRAPH.CENTER)

    set_cell_text(work.cell(1, 0), "1", align=WD_ALIGN_PARAGRAPH.CENTER)
    set_cell_text(work.cell(1, 1), "Hoàn thiện trang chủ và bộ sưu tập")
    set_cell_text(
        work.cell(1, 2),
        "Cập nhật trang chủ với hero banner, danh mục nổi bật, nhóm sản phẩm bán chạy và các điểm nhấn dịch vụ; đồng thời hoàn thiện trang bộ sưu tập với lưới sản phẩm, breadcrumb, thanh lọc theo danh mục/chất liệu/mức giá và các nút thao tác nhanh như xem chi tiết, thêm vào giỏ hàng và yêu thích.",
    )

    set_cell_text(work.cell(2, 0), "2", align=WD_ALIGN_PARAGRAPH.CENTER)
    set_cell_text(work.cell(2, 1), "Xây dựng trang chi tiết sản phẩm")
    set_cell_text(
        work.cell(2, 2),
        "Thiết kế trang chi tiết có hình ảnh lớn, ảnh thu nhỏ, hiển thị giá, trạng thái tồn kho, chọn kích cỡ, thêm vào giỏ hàng, thêm vào danh sách yêu thích và hai tab mô tả/thông tin chi tiết để người dùng xem nhanh toàn bộ dữ liệu sản phẩm.",
    )

    set_cell_text(work.cell(3, 0), "3", align=WD_ALIGN_PARAGRAPH.CENTER)
    set_cell_text(work.cell(3, 1), "Phát triển luồng tài khoản và mua hàng")
    set_cell_text(
        work.cell(3, 2),
        "Hoàn thiện các màn hình đăng nhập, đăng ký, hồ sơ cá nhân, đổi mật khẩu và địa chỉ giao hàng; đồng thời triển khai giỏ hàng theo session, cập nhật số lượng, xóa sản phẩm, trang thanh toán với thông tin giao hàng và phương thức thanh toán COD/chuyển khoản.",
    )

    set_cell_text(work.cell(4, 0), "4", align=WD_ALIGN_PARAGRAPH.CENTER)
    set_cell_text(work.cell(4, 1), "Bổ sung khu quản trị")
    set_cell_text(
        work.cell(4, 2),
        "Hoàn thiện dashboard quản trị và các màn hình CRUD cho sản phẩm, danh mục, khách hàng, đơn hàng, coupon, banner, blog, FAQ, báo cáo doanh thu và tồn kho; đồng thời thêm kiểm tra phân quyền cho tài khoản quản trị/nhân viên để bảo vệ các thao tác chỉnh sửa dữ liệu.",
    )

    add_text_paragraph(
        doc,
        "3. Kết quả đạt được:",
        style="Heading 3",
        alignment=WD_ALIGN_PARAGRAPH.LEFT,
        after=6,
        line=1.0,
        size=HEADING_SIZE,
        color=BLUE,
        bold=True,
    )
    add_text_paragraph(
        doc,
        "Đến cuối tuần 3, website PandoraWeb đã hình thành rõ luồng sử dụng chính của một cửa hàng trang sức trực tuyến: người dùng có thể vào trang chủ, duyệt bộ sưu tập, xem chi tiết, đưa sản phẩm vào giỏ hàng, đi tới thanh toán và nhận phản hồi đặt hàng thành công. Phần giao diện quản trị cũng đã có khung làm việc rõ ràng để quản lý sản phẩm, danh mục và các dữ liệu nghiệp vụ quan trọng.",
        style="Normal",
        alignment=WD_ALIGN_PARAGRAPH.JUSTIFY,
        after=6,
        line=1.5,
        first_line=0.5,
        size=BODY_SIZE,
        color=BLACK,
    )
    add_text_paragraph(
        doc,
        "Một số khu vực như bộ lọc sản phẩm, các khung mô tả chi tiết hoặc vài trang quản trị vẫn còn dùng dữ liệu mẫu ở một số vị trí, nhưng cấu trúc tổng thể đã sẵn sàng để thay thế dần bằng dữ liệu thật. So với tuần trước, hệ thống đã tiến thêm một bước lớn từ việc mới có nền tảng backend sang mức có thể chạy được một hành trình mua sắm tương đối hoàn chỉnh.",
        style="Normal",
        alignment=WD_ALIGN_PARAGRAPH.JUSTIFY,
        after=6,
        line=1.5,
        first_line=0.5,
        size=BODY_SIZE,
        color=BLACK,
    )

    add_text_paragraph(
        doc,
        "4. Khó khăn và hướng xử lý:",
        style="Heading 3",
        alignment=WD_ALIGN_PARAGRAPH.LEFT,
        after=6,
        line=1.0,
        size=HEADING_SIZE,
        color=BLUE,
        bold=True,
    )

    difficulties = doc.add_table(rows=4, cols=3)
    style_table(difficulties, [850, 2600, 5910])
    set_cell_text(difficulties.cell(0, 0), "STT", bold=True, align=WD_ALIGN_PARAGRAPH.CENTER)
    set_cell_text(difficulties.cell(0, 1), "Khó khăn", bold=True, align=WD_ALIGN_PARAGRAPH.CENTER)
    set_cell_text(difficulties.cell(0, 2), "Hướng xử lý", bold=True, align=WD_ALIGN_PARAGRAPH.CENTER)

    set_cell_text(difficulties.cell(1, 0), "1", align=WD_ALIGN_PARAGRAPH.CENTER)
    set_cell_text(difficulties.cell(1, 1), "Giao diện có nhiều thành phần tĩnh")
    set_cell_text(
        difficulties.cell(1, 2),
        "Ưu tiên giữ bố cục ổn định, sau đó thay dần các nội dung mẫu bằng dữ liệu lấy từ database hoặc API để tránh làm rối toàn bộ trang trong cùng một thời điểm.",
    )

    set_cell_text(difficulties.cell(2, 0), "2", align=WD_ALIGN_PARAGRAPH.CENTER)
    set_cell_text(difficulties.cell(2, 1), "Giỏ hàng và tồn kho phải đồng bộ")
    set_cell_text(
        difficulties.cell(2, 2),
        "Kiểm tra kỹ luồng thêm/xóa/cập nhật số lượng và trừ tồn kho sau khi đặt hàng, đồng thời dùng session và truy vấn cơ sở dữ liệu nhất quán để hạn chế sai lệch dữ liệu.",
    )

    set_cell_text(difficulties.cell(3, 0), "3", align=WD_ALIGN_PARAGRAPH.CENTER)
    set_cell_text(difficulties.cell(3, 1), "Quản lý phân quyền và trải nghiệm mobile")
    set_cell_text(
        difficulties.cell(3, 2),
        "Dùng kiểm tra quyền ở controller, giữ các màn hình quản trị theo mô hình modal/table rõ ràng, và tiếp tục tối ưu responsive để các trang chính hiển thị gọn trên điện thoại.",
    )

    add_text_paragraph(
        doc,
        "5. Kế hoạch tuần 4:",
        style="Heading 3",
        alignment=WD_ALIGN_PARAGRAPH.LEFT,
        after=6,
        line=1.0,
        size=HEADING_SIZE,
        color=BLUE,
        bold=True,
    )
    add_text_paragraph(
        doc,
        "Trong tuần 4, em sẽ tiếp tục hoàn thiện các phần còn dang dở như tìm kiếm, sắp xếp, phân trang và đồng bộ dữ liệu thật cho các bộ lọc trên trang sản phẩm.",
        style="Normal",
        alignment=WD_ALIGN_PARAGRAPH.JUSTIFY,
        after=4,
        line=1.5,
        first_line=0.5,
        size=BODY_SIZE,
        color=BLACK,
    )
    add_text_paragraph(
        doc,
        "Song song với đó, em sẽ hoàn thiện nốt các màn hình quản trị còn lại, rà soát lại validation của form đăng nhập - đăng ký - thanh toán và kiểm thử kỹ hơn các trường hợp biên như giỏ hàng rỗng, tồn kho thấp hoặc tài khoản chưa có địa chỉ mặc định.",
        style="Normal",
        alignment=WD_ALIGN_PARAGRAPH.JUSTIFY,
        after=4,
        line=1.5,
        first_line=0.5,
        size=BODY_SIZE,
        color=BLACK,
    )
    add_text_paragraph(
        doc,
        "Cuối cùng, em sẽ kiểm tra lại toàn bộ giao diện trên nhiều kích thước màn hình, chỉnh sửa các lỗi phát sinh và chuẩn bị dữ liệu minh họa để dự án bước sang giai đoạn hoàn thiện và báo cáo tổng kết.",
        style="Normal",
        alignment=WD_ALIGN_PARAGRAPH.JUSTIFY,
        after=6,
        line=1.5,
        first_line=0.5,
        size=BODY_SIZE,
        color=BLACK,
    )

    add_text_paragraph(
        doc,
        "6. Kết luận",
        style="Heading 3",
        alignment=WD_ALIGN_PARAGRAPH.LEFT,
        after=6,
        line=1.0,
        size=HEADING_SIZE,
        color=BLUE,
        bold=True,
    )
    add_text_paragraph(
        doc,
        "Tuần 3 là giai đoạn dự án PandoraWeb chuyển từ nền tảng kỹ thuật sang hình thái ứng dụng có thể sử dụng được theo đúng luồng thương mại điện tử. Các chức năng hiển thị sản phẩm, tài khoản, giỏ hàng, thanh toán và quản trị đã kết nối thành một hệ thống tương đối hoàn chỉnh, tạo nền tảng vững chắc để em tiếp tục tinh chỉnh dữ liệu, kiểm thử và hoàn thiện ở các tuần sau.",
        style="Normal",
        alignment=WD_ALIGN_PARAGRAPH.JUSTIFY,
        after=0,
        line=1.5,
        first_line=0.5,
        size=BODY_SIZE,
        color=BLACK,
    )

    doc.save(OUT_PATH)
    return OUT_PATH


if __name__ == "__main__":
    path = build_report()
    print("Saved week 3 report.")
