# -*- coding: utf-8 -*-
from pathlib import Path
import sys

from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

DOCS_SCRIPTS = Path(r"C:\Users\thean\.codex\plugins\cache\openai-primary-runtime\documents\26.623.12021\skills\documents\scripts")
if str(DOCS_SCRIPTS) not in sys.path:
    sys.path.append(str(DOCS_SCRIPTS))

try:
    from table_geometry import apply_table_geometry
except ImportError:
    def apply_table_geometry(table, widths_dxa, **kwargs):
        pass

ROOT = Path(__file__).resolve().parents[1]
OUT_PATH = ROOT / "Baocaotiendotuan4_NguyenTheAn_2400004657.docx"

FONT_NAME = "Times New Roman"
BODY_SIZE = 13
TITLE_SIZE = 16
HEADING_SIZE = 14
BLUE = RGBColor(15, 71, 97)
BLACK = RGBColor(0, 0, 0)

def set_run_font(run, *, size=None, color=None, bold=None, italic=None):
    run.font.name = FONT_NAME
    run._element.rPr.rFonts.set(qn("w:ascii"), FONT_NAME)
    run._element.rPr.rFonts.set(qn("w:hAnsi"), FONT_NAME)
    run._element.rPr.rFonts.set(qn("w:eastAsia"), FONT_NAME)
    if size is not None:
        run.font.size = Pt(size)
    if color is not None:
        run.font.color.rgb = color
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic

def set_style_font(style, *, size=None, color=None, bold=None, italic=None):
    style.font.name = FONT_NAME
    style.element.rPr.rFonts.set(qn("w:ascii"), FONT_NAME)
    style.element.rPr.rFonts.set(qn("w:hAnsi"), FONT_NAME)
    style.element.rPr.rFonts.set(qn("w:eastAsia"), FONT_NAME)
    if size is not None:
        style.font.size = Pt(size)
    if color is not None:
        style.font.color.rgb = color
    if bold is not None:
        style.font.bold = bold
    if italic is not None:
        style.font.italic = italic

def format_paragraph(paragraph, *, alignment=WD_ALIGN_PARAGRAPH.JUSTIFY, before=0, after=0, line=1.5, first_line=0):
    paragraph.alignment = alignment
    fmt = paragraph.paragraph_format
    fmt.space_before = Pt(before)
    fmt.space_after = Pt(after)
    fmt.line_spacing = line
    if first_line:
        fmt.first_line_indent = Inches(first_line)

def add_text_paragraph(doc, text, *, style="Normal", alignment=WD_ALIGN_PARAGRAPH.JUSTIFY, before=0, after=0, line=1.5, first_line=0, size=BODY_SIZE, color=BLACK, bold=False, italic=False):
    p = doc.add_paragraph(style=style)
    format_paragraph(p, alignment=alignment, before=before, after=after, line=line, first_line=first_line)
    run = p.add_run(text)
    set_run_font(run, size=size, color=color, bold=bold, italic=italic)
    return p

def clear_cell(cell):
    cell.text = ""
    for p in cell.paragraphs:
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(0)
        p.paragraph_format.line_spacing = 1.0

def set_cell_text(cell, text, *, bold=False, align=WD_ALIGN_PARAGRAPH.LEFT, size=BODY_SIZE):
    clear_cell(cell)
    p = cell.paragraphs[0]
    p.alignment = align
    fmt = p.paragraph_format
    fmt.space_before = Pt(0)
    fmt.space_after = Pt(0)
    fmt.line_spacing = 1.15
    run = p.add_run(text)
    set_run_font(run, size=size, bold=bold, color=BLACK)

def style_table(table, widths_dxa):
    apply_table_geometry(
        table,
        widths_dxa,
        table_width_dxa=sum(widths_dxa),
        indent_dxa=120,
        cell_margins_dxa={"top": 80, "bottom": 80, "start": 120, "end": 120},
    )
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    for row in table.rows:
        for cell in row.cells:
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.TOP
            for p in cell.paragraphs:
                p.paragraph_format.space_before = Pt(0)
                p.paragraph_format.space_after = Pt(0)
                p.paragraph_format.line_spacing = 1.15

def build_report():
    doc = Document()
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)

    normal = doc.styles["Normal"]
    set_style_font(normal, size=BODY_SIZE, color=BLACK)
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(0)
    normal.paragraph_format.line_spacing = 1.5

    h2 = doc.styles["Heading 2"]
    set_style_font(h2, size=TITLE_SIZE, color=BLUE, bold=True)
    h2.paragraph_format.space_before = Pt(0)
    h2.paragraph_format.space_after = Pt(6)
    h2.paragraph_format.line_spacing = 1.0

    h3 = doc.styles["Heading 3"]
    set_style_font(h3, size=HEADING_SIZE, color=BLUE, bold=True)
    h3.paragraph_format.space_before = Pt(12)
    h3.paragraph_format.space_after = Pt(6)
    h3.paragraph_format.line_spacing = 1.0

    doc.core_properties.title = "Báo cáo tiến độ tuần 4"
    doc.core_properties.author = "Nguyễn Thế An"
    doc.core_properties.subject = "PandoraWeb"

    add_text_paragraph(doc, "BÁO CÁO TIẾN ĐỘ TUẦN 4", style="Heading 2", alignment=WD_ALIGN_PARAGRAPH.LEFT, after=6, line=1.0, size=TITLE_SIZE, color=BLUE, bold=True)
    add_text_paragraph(doc, "Đề tài: Xây dựng hệ thống thương mại điện tử quản lý và kinh doanh trang sức trực tuyến", style="Normal", alignment=WD_ALIGN_PARAGRAPH.JUSTIFY, after=8, line=1.5, size=BODY_SIZE, color=BLACK)

    meta = doc.add_table(rows=3, cols=2)
    style_table(meta, [2940, 6420])
    set_cell_text(meta.cell(0, 0), "Sinh viên thực hiện", bold=True)
    set_cell_text(meta.cell(0, 1), "Nguyễn Thế An")
    set_cell_text(meta.cell(1, 0), "MSSV", bold=True)
    set_cell_text(meta.cell(1, 1), "2400004657")
    set_cell_text(meta.cell(2, 0), "Thời gian thực hiện", bold=True)
    set_cell_text(meta.cell(2, 1), "Tuần 4")

    add_text_paragraph(doc, "1. Tổng quan tiến độ tuần 4:", style="Heading 3", alignment=WD_ALIGN_PARAGRAPH.LEFT, after=6, line=1.0, size=HEADING_SIZE, color=BLUE, bold=True)
    add_text_paragraph(doc, "Trong tuần 4, dự án PandoraWeb tập trung vào việc hoàn thiện quy trình thanh toán (Checkout) và xử lý nghiệp vụ đơn hàng dưới Database. Mục tiêu là giúp quy trình mua hàng diễn ra trơn tru cho cả người dùng đã có tài khoản lẫn khách vãng lai, đồng thời đảm bảo dữ liệu đơn hàng và tồn kho được quản lý chính xác, minh bạch.", style="Normal", alignment=WD_ALIGN_PARAGRAPH.JUSTIFY, after=6, line=1.5, first_line=0.5, size=BODY_SIZE, color=BLACK)

    add_text_paragraph(doc, "2. Công việc đã thực hiện trong tuần 4:", style="Heading 3", alignment=WD_ALIGN_PARAGRAPH.LEFT, after=6, line=1.0, size=HEADING_SIZE, color=BLUE, bold=True)
    add_text_paragraph(doc, "Nội dung công việc tuần 4 được triển khai bám sát 4 yêu cầu nghiệp vụ cốt lõi của phân hệ đặt hàng và thanh toán:", style="Normal", alignment=WD_ALIGN_PARAGRAPH.JUSTIFY, after=6, line=1.5, first_line=0.5, size=BODY_SIZE, color=BLACK)

    work = doc.add_table(rows=5, cols=3)
    style_table(work, [850, 2500, 6010])
    set_cell_text(work.cell(0, 0), "STT", bold=True, align=WD_ALIGN_PARAGRAPH.CENTER)
    set_cell_text(work.cell(0, 1), "Hạng mục", bold=True, align=WD_ALIGN_PARAGRAPH.CENTER)
    set_cell_text(work.cell(0, 2), "Nội dung thực hiện", bold=True, align=WD_ALIGN_PARAGRAPH.CENTER)

    set_cell_text(work.cell(1, 0), "1", align=WD_ALIGN_PARAGRAPH.CENTER)
    set_cell_text(work.cell(1, 1), "Tích hợp trang Thanh toán nhận thông tin giao hàng")
    set_cell_text(work.cell(1, 2), "Xây dựng giao diện và Controller thu thập thông tin nhận hàng của khách (Họ tên, SĐT, Email, Địa chỉ chi tiết, Ghi chú). Tự động điền sẵn thông tin nếu khách đã đăng nhập. Hỗ trợ phương thức thanh toán COD và Chuyển khoản ngân hàng (QR Code).")

    set_cell_text(work.cell(2, 0), "2", align=WD_ALIGN_PARAGRAPH.CENTER)
    set_cell_text(work.cell(2, 1), "Tự động tạo Khách hàng vãng lai (Guest Checkout)")
    set_cell_text(work.cell(2, 2), "Xây dựng luồng xử lý kiểm tra email khi chốt đơn. Nếu người mua chưa có tài khoản, hệ thống tự động khởi tạo hồ sơ Khách hàng vãng lai (Guest) trong Database với mật khẩu mặc định, đồng thời duy trì Session đăng nhập ngay để người dùng xem lại đơn hàng.")

    set_cell_text(work.cell(3, 0), "3", align=WD_ALIGN_PARAGRAPH.CENTER)
    set_cell_text(work.cell(3, 1), "Chuyển dữ liệu Giỏ hàng (Session) lưu vĩnh viễn xuống Database")
    set_cell_text(work.cell(3, 2), "Viết hàm chuyển đổi toàn bộ thông tin giỏ hàng tạm thời từ Session thành các bản ghi vĩnh viễn trong 2 bảng Orders và OrderItems. Sau khi lưu thành công, hệ thống xóa sạch giỏ hàng tạm và tự động gửi Email hóa đơn xác nhận về Gmail người mua.")

    set_cell_text(work.cell(4, 0), "4", align=WD_ALIGN_PARAGRAPH.CENTER)
    set_cell_text(work.cell(4, 1), "Thuật toán Tự động trừ tồn kho khi đặt hàng thành công")
    set_cell_text(work.cell(4, 2), "Áp dụng Database Transaction (BeginTransaction) kiểm tra số lượng tồn kho từng biến thể sản phẩm. Nếu đủ số lượng, hệ thống thực hiện trừ tồn kho (Stock -= Quantity) và chốt giao dịch. Nếu không đủ, tự động Rollback hoàn tác và thông báo lỗi rõ ràng.")

    add_text_paragraph(doc, "3. Kết quả đạt được:", style="Heading 3", alignment=WD_ALIGN_PARAGRAPH.LEFT, after=6, line=1.0, size=HEADING_SIZE, color=BLUE, bold=True)
    add_text_paragraph(doc, "Đến cuối tuần 4, luồng mua sắm trực tuyến trên PandoraWeb đã hoạt động hoàn chỉnh từ bước chọn hàng, xem giỏ hàng, điền thông tin thanh toán, tạo đơn hàng vĩnh viễn trong Database, trừ tồn kho chính xác cho đến bước gửi email xác nhận cho khách hàng. Hệ thống hỗ trợ tốt cả hai đối tượng: khách đã đăng nhập và khách vãng lai mua nhanh.", style="Normal", alignment=WD_ALIGN_PARAGRAPH.JUSTIFY, after=6, line=1.5, first_line=0.5, size=BODY_SIZE, color=BLACK)

    add_text_paragraph(doc, "4. Khó khăn và hướng xử lý:", style="Heading 3", alignment=WD_ALIGN_PARAGRAPH.LEFT, after=6, line=1.0, size=HEADING_SIZE, color=BLUE, bold=True)

    diff = doc.add_table(rows=3, cols=3)
    style_table(diff, [850, 2600, 5910])
    set_cell_text(diff.cell(0, 0), "STT", bold=True, align=WD_ALIGN_PARAGRAPH.CENTER)
    set_cell_text(diff.cell(0, 1), "Khó khăn", bold=True, align=WD_ALIGN_PARAGRAPH.CENTER)
    set_cell_text(diff.cell(0, 2), "Hướng xử lý", bold=True, align=WD_ALIGN_PARAGRAPH.CENTER)

    set_cell_text(diff.cell(1, 0), "1", align=WD_ALIGN_PARAGRAPH.CENTER)
    set_cell_text(diff.cell(1, 1), "Tranh chấp dữ liệu tồn kho khi nhiều người mua cùng lúc")
    set_cell_text(diff.cell(1, 2), "Sử dụng Database Transaction để khóa và kiểm tra số lượng tồn kho ngay trước khi trừ, đảm bảo không bị âm kho hoặc quá tải đơn.")

    set_cell_text(diff.cell(2, 0), "2", align=WD_ALIGN_PARAGRAPH.CENTER)
    set_cell_text(diff.cell(2, 1), "Xử lý trải nghiệm đặt hàng cho khách chưa đăng nhập")
    set_cell_text(diff.cell(2, 2), "Khởi tạo tự động tài khoản Guest dựa trên Email đăng ký nhận hàng, giúp lưu vết đơn hàng mà không làm gián đoạn trải nghiệm người dùng.")

    add_text_paragraph(doc, "5. Kế hoạch tuần 5:", style="Heading 3", alignment=WD_ALIGN_PARAGRAPH.LEFT, after=6, line=1.0, size=HEADING_SIZE, color=BLUE, bold=True)
    add_text_paragraph(doc, "Trong tuần 5, em sẽ tiếp tục triển khai các hạng mục bảo mật và hoàn thiện quản lý tài khoản người dùng:", style="Normal", alignment=WD_ALIGN_PARAGRAPH.JUSTIFY, after=4, line=1.5, first_line=0.5, size=BODY_SIZE, color=BLACK)
    
    add_text_paragraph(doc, "• Biến form Đăng ký và Đăng nhập tĩnh thành form động, thực hiện mã hóa mật khẩu an toàn xuống Database.", style="Normal", alignment=WD_ALIGN_PARAGRAPH.JUSTIFY, after=3, line=1.5, first_line=0.5, size=BODY_SIZE, color=BLACK)
    add_text_paragraph(doc, "• Áp dụng bộ lọc bảo vệ ([AdminAuthorize]) để ngăn chặn triệt để người dùng không có quyền truy cập vào các trang Quản trị Admin.", style="Normal", alignment=WD_ALIGN_PARAGRAPH.JUSTIFY, after=3, line=1.5, first_line=0.5, size=BODY_SIZE, color=BLACK)
    add_text_paragraph(doc, "• Hoàn thiện trang Hồ sơ cá nhân: Cho phép người dùng đổi mật khẩu, cập nhật thông tin, xem lại danh sách đơn hàng đã đặt và theo dõi chi tiết trạng thái giao hàng.", style="Normal", alignment=WD_ALIGN_PARAGRAPH.JUSTIFY, after=6, line=1.5, first_line=0.5, size=BODY_SIZE, color=BLACK)

    add_text_paragraph(doc, "6. Kết luận:", style="Heading 3", alignment=WD_ALIGN_PARAGRAPH.LEFT, after=6, line=1.0, size=HEADING_SIZE, color=BLUE, bold=True)
    add_text_paragraph(doc, "Tuần 4 đã giải quyết xong toàn bộ phần cốt lõi của nghiệp vụ Đặt hàng - Thanh toán - Trừ tồn kho - Gửi hóa đơn Email. Đây là bước tiến quan trọng giúp PandoraWeb đạt được tiêu chuẩn của một trang web thương mại điện tử thực thụ.", style="Normal", alignment=WD_ALIGN_PARAGRAPH.JUSTIFY, after=0, line=1.5, first_line=0.5, size=BODY_SIZE, color=BLACK)

    doc.save(OUT_PATH)
    return OUT_PATH

if __name__ == "__main__":
    path = build_report()
    print(f"Saved week 4 report to {path}")
