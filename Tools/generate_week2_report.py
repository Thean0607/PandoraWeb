import sys
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

DOCS_SCRIPTS = Path(r"C:\Users\thean\.codex\plugins\cache\openai-primary-runtime\documents\26.623.12021\skills\documents\scripts")
if str(DOCS_SCRIPTS) not in sys.path:
    sys.path.append(str(DOCS_SCRIPTS))

from table_geometry import apply_table_geometry


OUT_PATH = Path(r"C:\Users\thean\Desktop\Đồ án cơ sở\PandoraWeb\Baocaotiendotuan2_NguyenTheAn_2400004657.docx")

FONT_NAME = "Times New Roman"
BODY_SIZE = 13
TITLE_SIZE = 16
HEADING_SIZE = 14
CAPTION_SIZE = 9
BLUE = RGBColor(15, 71, 97)
BLACK = RGBColor(0, 0, 0)


def set_run_font(run, *, size=None, color=None, bold=None, italic=None):
    run.font.name = FONT_NAME
    run._element.rPr.rFonts.set(qn("w:ascii"), FONT_NAME)
    run._element.rPr.rFonts.set(qn("w:hAnsi"), FONT_NAME)
    run._element.rPr.rFonts.set(qn("w:eastAsia"), FONT_NAME)
    if size is not None:
        run.font.size = Pt(size)
    if color is not None:
        run.font.color.rgb = color
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic


def set_style_font(style, *, size=None, color=None, bold=None, italic=None):
    style.font.name = FONT_NAME
    style.element.rPr.rFonts.set(qn("w:ascii"), FONT_NAME)
    style.element.rPr.rFonts.set(qn("w:hAnsi"), FONT_NAME)
    style.element.rPr.rFonts.set(qn("w:eastAsia"), FONT_NAME)
    if size is not None:
        style.font.size = Pt(size)
    if color is not None:
        style.font.color.rgb = color
    if bold is not None:
        style.font.bold = bold
    if italic is not None:
        style.font.italic = italic


def format_paragraph(paragraph, *, alignment=WD_ALIGN_PARAGRAPH.JUSTIFY, before=0, after=0, line=1.5, first_line=0):
    paragraph.alignment = alignment
    fmt = paragraph.paragraph_format
    fmt.space_before = Pt(before)
    fmt.space_after = Pt(after)
    fmt.line_spacing = line
    if first_line:
        fmt.first_line_indent = Inches(first_line)


def add_text_paragraph(doc, text, *, style="Normal", alignment=WD_ALIGN_PARAGRAPH.JUSTIFY, before=0, after=0, line=1.5, first_line=0, size=BODY_SIZE, color=BLACK, bold=False, italic=False):
    p = doc.add_paragraph(style=style)
    format_paragraph(p, alignment=alignment, before=before, after=after, line=line, first_line=first_line)
    run = p.add_run(text)
    set_run_font(run, size=size, color=color, bold=bold, italic=italic)
    return p


def clear_cell(cell):
    cell.text = ""
    for p in cell.paragraphs:
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(0)
        p.paragraph_format.line_spacing = 1.0


def set_cell_text(cell, text, *, bold=False, align=WD_ALIGN_PARAGRAPH.LEFT, size=BODY_SIZE):
    clear_cell(cell)
    p = cell.paragraphs[0]
    p.alignment = align
    fmt = p.paragraph_format
    fmt.space_before = Pt(0)
    fmt.space_after = Pt(0)
    fmt.line_spacing = 1.15
    run = p.add_run(text)
    set_run_font(run, size=size, bold=bold, color=BLACK)


def style_table(table, widths_dxa):
    apply_table_geometry(
        table,
        widths_dxa,
        table_width_dxa=sum(widths_dxa),
        indent_dxa=120,
        cell_margins_dxa={"top": 80, "bottom": 80, "start": 120, "end": 120},
    )
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    for row in table.rows:
        for cell in row.cells:
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.TOP
            for p in cell.paragraphs:
                p.paragraph_format.space_before = Pt(0)
                p.paragraph_format.space_after = Pt(0)
                p.paragraph_format.line_spacing = 1.15


def main():
    doc = Document()
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    # Base styles
    normal = doc.styles["Normal"]
    set_style_font(normal, size=BODY_SIZE, color=BLACK)
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(0)
    normal.paragraph_format.line_spacing = 1.5

    h2 = doc.styles["Heading 2"]
    set_style_font(h2, size=TITLE_SIZE, color=BLUE, bold=True)
    h2.paragraph_format.space_before = Pt(0)
    h2.paragraph_format.space_after = Pt(6)
    h2.paragraph_format.line_spacing = 1.0

    h3 = doc.styles["Heading 3"]
    set_style_font(h3, size=HEADING_SIZE, color=BLUE, bold=True)
    h3.paragraph_format.space_before = Pt(12)
    h3.paragraph_format.space_after = Pt(6)
    h3.paragraph_format.line_spacing = 1.0

    caption = doc.styles["Caption"]
    set_style_font(caption, size=CAPTION_SIZE, color=RGBColor(14, 40, 65), italic=False)

    doc.core_properties.title = "Báo cáo tiến độ tuần 2"
    doc.core_properties.author = "Nguyễn Thế An"
    doc.core_properties.subject = "PandoraWeb"

    # Title block
    add_text_paragraph(
        doc,
        "BÁO CÁO TIẾN ĐỘ TUẦN 2",
        style="Heading 2",
        alignment=WD_ALIGN_PARAGRAPH.LEFT,
        after=6,
        line=1.0,
        size=TITLE_SIZE,
        color=BLUE,
        bold=True,
    )
    add_text_paragraph(
        doc,
        "Đề tài: Xây dựng hệ thống thương mại điện tử quản lý và kinh doanh trang sức trực tuyến",
        style="Normal",
        alignment=WD_ALIGN_PARAGRAPH.JUSTIFY,
        after=8,
        line=1.5,
        first_line=0,
        size=BODY_SIZE,
        color=BLACK,
    )

    # Metadata table
    meta = doc.add_table(rows=3, cols=2)
    style_table(meta, [2940, 6420])
    meta.rows[0].cells[0].text = ""
    meta.rows[0].cells[1].text = ""
    set_cell_text(meta.cell(0, 0), "Sinh viên thực hiện", bold=True)
    set_cell_text(meta.cell(0, 1), "Nguyễn Thế An")
    set_cell_text(meta.cell(1, 0), "MSSV", bold=True)
    set_cell_text(meta.cell(1, 1), "2400004657")
    set_cell_text(meta.cell(2, 0), "Thời gian thực hiện", bold=True)
    set_cell_text(meta.cell(2, 1), "Tuần 2")

    add_text_paragraph(
        doc,
        "1. Tổng quan tiến độ tuần 2:",
        style="Heading 3",
        alignment=WD_ALIGN_PARAGRAPH.LEFT,
        after=6,
        line=1.0,
        size=HEADING_SIZE,
        color=BLUE,
        bold=True,
    )
    add_text_paragraph(
        doc,
        "Sau khi hoàn thành khảo sát yêu cầu và dựng giao diện tĩnh ở tuần 1, tuần 2 của dự án được chuyển sang giai đoạn xây dựng nền tảng backend. Trong tuần này, em tập trung vào việc chuẩn hóa mô hình dữ liệu, thiết lập DbContext, cấu hình kết nối SQL Server và viết các controller/action để hệ thống PandoraWeb bắt đầu hiển thị dữ liệu động trên khung ASP.NET MVC.",
        style="Normal",
        alignment=WD_ALIGN_PARAGRAPH.JUSTIFY,
        after=6,
        line=1.5,
        first_line=0.5,
        size=BODY_SIZE,
        color=BLACK,
    )

    add_text_paragraph(
        doc,
        "2. Công việc đã thực hiện trong tuần 2:",
        style="Heading 3",
        alignment=WD_ALIGN_PARAGRAPH.LEFT,
        after=6,
        line=1.0,
        size=HEADING_SIZE,
        color=BLUE,
        bold=True,
    )
    add_text_paragraph(
        doc,
        "Các đầu việc trong tuần này được triển khai theo hướng ưu tiên lớp dữ liệu trước, sau đó mới đồng bộ sang lớp xử lý và giao diện để bảo đảm cấu trúc dự án ổn định ngay từ giai đoạn đầu.",
        style="Normal",
        alignment=WD_ALIGN_PARAGRAPH.JUSTIFY,
        after=6,
        line=1.5,
        first_line=0.5,
        size=BODY_SIZE,
        color=BLACK,
    )

    work = doc.add_table(rows=5, cols=3)
    style_table(work, [850, 2500, 6010])
    set_cell_text(work.cell(0, 0), "STT", bold=True, align=WD_ALIGN_PARAGRAPH.CENTER)
    set_cell_text(work.cell(0, 1), "Hạng mục", bold=True, align=WD_ALIGN_PARAGRAPH.CENTER)
    set_cell_text(work.cell(0, 2), "Nội dung thực hiện", bold=True, align=WD_ALIGN_PARAGRAPH.CENTER)
    set_cell_text(work.cell(1, 0), "1", align=WD_ALIGN_PARAGRAPH.CENTER)
    set_cell_text(
        work.cell(1, 1),
        "Thiết kế và chuẩn hóa cơ sở dữ liệu",
        bold=False,
        align=WD_ALIGN_PARAGRAPH.LEFT,
    )
    set_cell_text(
        work.cell(1, 2),
        "Xác định các thực thể chính gồm Role, Employee, Customer, Category, Product, Order và OrderItem; đồng thời khai báo các quan hệ giữa chúng trong lớp PandoraDbContext để phục vụ cho việc lưu trữ và truy xuất dữ liệu.",
    )
    set_cell_text(work.cell(2, 0), "2", align=WD_ALIGN_PARAGRAPH.CENTER)
    set_cell_text(
        work.cell(2, 1),
        "Xây dựng kiến trúc MVC",
        bold=False,
        align=WD_ALIGN_PARAGRAPH.LEFT,
    )
    set_cell_text(
        work.cell(2, 2),
        "Hoàn thiện các controller chính: Home, Product, Order, Account và Admin; cấu hình RouteConfig, BundleConfig và Global.asax để ứng dụng khởi tạo đúng luồng và nạp tài nguyên đồng bộ.",
    )
    set_cell_text(work.cell(3, 0), "3", align=WD_ALIGN_PARAGRAPH.CENTER)
    set_cell_text(
        work.cell(3, 1),
        "Liên kết giao diện với dữ liệu động",
        bold=False,
        align=WD_ALIGN_PARAGRAPH.LEFT,
    )
    set_cell_text(
        work.cell(3, 2),
        "Trang chủ đã lấy 4 sản phẩm nổi bật từ database; trang danh mục và trang chi tiết sản phẩm đã sử dụng Include(p => p.Category) để hiển thị đúng thông tin liên quan.",
    )
    set_cell_text(work.cell(4, 0), "4", align=WD_ALIGN_PARAGRAPH.CENTER)
    set_cell_text(
        work.cell(4, 1),
        "Hoàn thiện các trang chức năng",
        bold=False,
        align=WD_ALIGN_PARAGRAPH.LEFT,
    )
    set_cell_text(
        work.cell(4, 2),
        "Cập nhật và đồng bộ các view quan trọng như giỏ hàng, thanh toán, đăng nhập, đăng ký, hồ sơ người dùng và khu vực quản trị nhằm chuẩn bị cho các xử lý nghiệp vụ tiếp theo.",
    )

    add_text_paragraph(
        doc,
        "3. Kết quả đạt được:",
        style="Heading 3",
        alignment=WD_ALIGN_PARAGRAPH.LEFT,
        after=6,
        line=1.0,
        size=HEADING_SIZE,
        color=BLUE,
        bold=True,
    )
    add_text_paragraph(
        doc,
        "Hệ thống đã hình thành rõ ràng 3 lớp chính: dữ liệu, xử lý và giao diện. Ở tầng dữ liệu, các bảng quan trọng của dự án đã được khai báo đầy đủ; ở tầng xử lý, các action đã sẵn sàng trả dữ liệu cho giao diện; ở tầng hiển thị, layout chung và các trang chức năng đã được tổ chức lại để có thể mở rộng dễ dàng trong các tuần sau.",
        style="Normal",
        alignment=WD_ALIGN_PARAGRAPH.JUSTIFY,
        after=6,
        line=1.5,
        first_line=0.5,
        size=BODY_SIZE,
        color=BLACK,
    )
    add_text_paragraph(
        doc,
        "Đặc biệt, một số luồng cơ bản như xem sản phẩm, xem danh mục và mở chi tiết sản phẩm đã chuyển từ giao diện tĩnh sang có khả năng hiển thị dữ liệu thật. Đây là bước quan trọng giúp dự án tiến gần hơn tới mô hình thương mại điện tử hoàn chỉnh thay vì chỉ dừng ở phần demo giao diện.",
        style="Normal",
        alignment=WD_ALIGN_PARAGRAPH.JUSTIFY,
        after=6,
        line=1.5,
        first_line=0.5,
        size=BODY_SIZE,
        color=BLACK,
    )

    add_text_paragraph(
        doc,
        "4. Khó khăn và hướng xử lý:",
        style="Heading 3",
        alignment=WD_ALIGN_PARAGRAPH.LEFT,
        after=6,
        line=1.0,
        size=HEADING_SIZE,
        color=BLUE,
        bold=True,
    )
    difficulties = doc.add_table(rows=3, cols=3)
    style_table(difficulties, [850, 2600, 5910])
    set_cell_text(difficulties.cell(0, 0), "STT", bold=True, align=WD_ALIGN_PARAGRAPH.CENTER)
    set_cell_text(difficulties.cell(0, 1), "Khó khăn", bold=True, align=WD_ALIGN_PARAGRAPH.CENTER)
    set_cell_text(difficulties.cell(0, 2), "Hướng xử lý", bold=True, align=WD_ALIGN_PARAGRAPH.CENTER)
    set_cell_text(difficulties.cell(1, 0), "1", align=WD_ALIGN_PARAGRAPH.CENTER)
    set_cell_text(difficulties.cell(1, 1), "Dữ liệu thật chưa hoàn chỉnh")
    set_cell_text(
        difficulties.cell(1, 2),
        "Tạm thời sử dụng dữ liệu mẫu và truy vấn đơn giản để kiểm tra giao diện; đồng thời tiếp tục chuẩn bị seed data và kế hoạch nhập dữ liệu ban đầu cho sản phẩm, danh mục và tài khoản.",
    )
    set_cell_text(difficulties.cell(2, 0), "2", align=WD_ALIGN_PARAGRAPH.CENTER)
    set_cell_text(difficulties.cell(2, 1), "Một số chức năng nghiệp vụ còn là khung giao diện")
    set_cell_text(
        difficulties.cell(2, 2),
        "Ưu tiên hoàn thiện kiến trúc controller/view trước, sau đó bổ sung logic cho giỏ hàng, thanh toán, đăng nhập và quản trị để tránh làm rối cấu trúc code.",
    )

    add_text_paragraph(
        doc,
        "5. Kế hoạch tuần 3:",
        style="Heading 3",
        alignment=WD_ALIGN_PARAGRAPH.LEFT,
        after=6,
        line=1.0,
        size=HEADING_SIZE,
        color=BLUE,
        bold=True,
    )
    add_text_paragraph(
        doc,
        "Trong tuần 3, dự án sẽ tập trung vào việc hoàn thiện dữ liệu đầu vào và bắt đầu hiện thực các thao tác CRUD cho sản phẩm, danh mục, nhân viên và khách hàng.",
        style="Normal",
        alignment=WD_ALIGN_PARAGRAPH.JUSTIFY,
        after=4,
        line=1.5,
        first_line=0.5,
        size=BODY_SIZE,
        color=BLACK,
    )
    add_text_paragraph(
        doc,
        "Song song với đó, em sẽ tiếp tục triển khai chức năng đăng nhập/đăng ký, xử lý giỏ hàng theo session hoặc database và chuẩn hóa luồng thanh toán để dữ liệu được lưu xuyên suốt.",
        style="Normal",
        alignment=WD_ALIGN_PARAGRAPH.JUSTIFY,
        after=4,
        line=1.5,
        first_line=0.5,
        size=BODY_SIZE,
        color=BLACK,
    )
    add_text_paragraph(
        doc,
        "Cuối cùng, em sẽ kiểm thử lại toàn bộ các trang đã kết nối dữ liệu động, tối ưu hiển thị trên nhiều kích thước màn hình và khắc phục các lỗi phát sinh trước khi bước sang giai đoạn hoàn thiện chức năng nâng cao.",
        style="Normal",
        alignment=WD_ALIGN_PARAGRAPH.JUSTIFY,
        after=6,
        line=1.5,
        first_line=0.5,
        size=BODY_SIZE,
        color=BLACK,
    )

    add_text_paragraph(
        doc,
        "6. Kết luận",
        style="Heading 3",
        alignment=WD_ALIGN_PARAGRAPH.LEFT,
        after=6,
        line=1.0,
        size=HEADING_SIZE,
        color=BLUE,
        bold=True,
    )
    add_text_paragraph(
        doc,
        "Tuần 2 đánh dấu bước chuyển quan trọng của dự án PandoraWeb từ giai đoạn dựng giao diện sang giai đoạn xây dựng nền tảng backend. Mặc dù các chức năng nghiệp vụ phức tạp vẫn cần tiếp tục hoàn thiện, những kết quả đạt được trong tuần này đã tạo ra bộ khung MVC, mô hình dữ liệu và lớp giao diện đủ vững để dự án tiếp tục phát triển ổn định trong các tuần tiếp theo.",
        style="Normal",
        alignment=WD_ALIGN_PARAGRAPH.JUSTIFY,
        after=0,
        line=1.5,
        first_line=0.5,
        size=BODY_SIZE,
        color=BLACK,
    )

    doc.save(OUT_PATH)
    print("Saved week 2 report.")


if __name__ == "__main__":
    main()
